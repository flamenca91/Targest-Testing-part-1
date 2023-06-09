# This program will search for red colored tags to get the information it needs
#
# Instructions on how to use the program:
# 1. run program
# 2. click on button "Choose document" and choose a document
# 3. Do step 2 as many times as you want, to uadd as many documents as you like
# 4. After you are done choosing documents, click on the GenerateReport button
# 5. Then you can click on the "open generated report" button, which will automatically
# open up your word document report created from your documents
# 6. When you are done, click "End Program"

# from debug import debug
import logging
# import pdb
import docx
from docx import Document
from docx.shared import RGBColor
from docx.shared import Inches
import tkinter as tk
from tkinter import *
from tkinter import filedialog
from typing import Tuple

from tkinter import scrolledtext
import re
import copy
import time

# This libraries are for opening word document automatically
#import os
#import platform
#import subprocess

# This library is for opening excel document automatically
#import xlwings as xw
#import pandas as pd
#import matplotlib.pyplot as plt



# Set up the logger for catching errors
logging.basicConfig(level=logging.ERROR,
                    format='%(asctime)s - %(levelname)s - %(message)s')

logger = logging.getLogger(__name__)

global doNothing
doNothing = 0


global dicts2Copy # This will hold the dicts2 content in all documents
dicts2Copy = {}

global parents2Copy # parents2 list copy
parents2Copy = []

global filtered_L # Will store the ones without a child tag
filtered_L = []

global filtered_LCopy
filtered_LCopy = []

global fullText2Copy
fullText2Copy = []

global parents2 #list of parent tags or child tags
parents2 = []

# creates a dict for parent and child tags
global dicts
dicts = {}

global OrphanChild2
OrphanChild2 = []
global OrphanChild2Copy
orphanChildren2Copy = []

global dicts10
dicts10 = {}
global dicts3
dicts3 = {}  # will hold parentTag and text, Orphan tags
global dicts2
dicts2 = {}  # will hold parentTag and text
global orphanDicts
orphanDicts = {}  # orphan dictionary

global parents9
parents9 = []

# declaring different lists that will be used to store, tags and sentences
global parentTags
parentTags = []
global parent  # This will be used to store everything
parent = []
global child # Used to Store child tags
child = [] 
global noChild # Used to Store parentTags with no child
noChild = []  
global withChild # Used to Store parentTags with child tag
withChild = [] 
global parents # Will be used for future function
parents = [] 

global orphanTagText
orphanTagText = []  # Will be used to hold text of orphanChildTags


# reads the text in the document and use the getcoloredTXT function to get the colored text
def readtxt(filename, color: Tuple[int, int, int]):
    try:
        doc = docx.Document(filename)
        text10 = ""
        fullText = []
        new = []
        global everything
        everything = []  # list of tags and text

        for para in doc.paragraphs:
            # Getting the colored words from the doc
            if (getcoloredTxt(para.runs, color)):
                # Concatenating list of runs between the colored text to single a string
                sentence = "".join(r.text for r in para.runs)
                if len(sentence) > 5:  # this chcekcts if sentence has atleast 5 characters
                    fullText.append(sentence)
                #print(sentence) # Prints everything in the terminal
                everything.append(sentence)
                text10 = sentence
                parent.append("".join(r.text for r in para.runs))

        #print(fullText)
        global hasChild # Will store the ones with a child tag
        global fullText2 # will store everything found
        global children

        global orphanss
        orphanss = []

        global orphanChildren2 # Will store the orphan child tags for orphanReport
        orphanChildren2 = []
        # Finds the lines without a parentTag
        filtered_L = [value for value in fullText if "[" not in value]

        filtered_L = [s.replace(": ", ":") for s in filtered_L]
        # Finds the lines with a parentTag
        filtered_LCopy.extend(filtered_L)
        hasChild = [value for value in fullText if "[" in value]
        # will store everything found
        fullText2 = [value for value in fullText]
        fullText2 = [s.replace(": ", ":") for s in fullText2]
        fullText2 = [s.replace("[ ", "[") for s in fullText2]
        fullText2 = [s.replace("] ", "]") for s in fullText2]
        fullText2Copy.extend(fullText2)
        
        return fullText, filtered_L, hasChild, filtered_LCopy, fullText2Copy, fullText2

    except Exception as e:
    # Log an error message
        logging.error('readtxt(): ERROR', exc_info=True)


def getcoloredTxt(runs, color): 
    coloredWords, word = [], ""
    try:
        for run in runs:
            if run.font.color.rgb == RGBColor(*color):
                word += str(run.text) # Saves everything found

            elif word != "":  # This will find the parentTags
                coloredWords.append(word)
                parentTags.append(word)
                parents.append(word)
                word = ""

        if word != "":  # This will find the parentTags
            coloredWords.append(word + "\n")
            # word = removeAfter(word)
            child.append(word)
            withChild.append(word)

            

    except Exception as e:
        logging.error('getColoredText(): ERROR', e)
    else:
        # Log a success message
        logging.info('getColoredText(): PASS')

    return coloredWords # returns everything found


def generateReport(): #Will generate the report for tags
    try:
        global filepath
        global filepath2
        # create a similar method for opening a folder
        filepath = filedialog.askopenfilename(initialdir="/",
                                            title="",
                                            filetypes = (("text file", "*.txt"),
                                                        ("all files","*.*")))
        file = open(filepath,'r')
        #print(filepath)
        file.close()
        # Will store the filepath to the document as a string
        filepath2 = str(filepath)
        a = (filepath2)
        with open(a) as file_in:
            lines = []
            for line in file_in:
                lines.append(line)
        for line2 in lines:
            print(line2)
            line3 = str(line2)
            line4 = line3.replace('\\', '/')
            line5 = line4.replace('"', '')
            line6 = line5.replace("\n", "")
            print(line6)
            fullText = readtxt(filename=line6,
                            color=(255, 0, 0))
            print(line4)

            #filtered_L = readtxt(filename=filepath2, #For future use
            #                   color=(255, 0, 0))
            fullText10 = str(fullText)
            s = ''.join(fullText10)
            w = (s.replace (']', ']\n\n'))
            #paragraph = report3.add_paragraph()

            #paragraph2 = orphanReport.add_paragraph()
            filepath3 = str(line4.rsplit('/', 1)[-1]) # change filepath to something.docx
            filepath3 = filepath3.split('.', 1)[0] # removes .docx of the file name
            print(filepath3 + " added to the report")
            nameOfDoc = (filepath3 + " added to the report\n")
    
            #runner = paragraph.add_run("\n" + "Document Name: " + filepath3 + "\n")
            #runner.bold = True  # makes the header bold



            # w will be used in the future
            w = (w.replace ('([', ''))
            w = (w.replace (',', ''))
            w = (w.replace ('' '', ''))

            # creates a table for report 1
            #table = report3.add_table(rows=1, cols=2)






            # Now save the document to a location
            #report3.save('report.docx')


            #orphanReport.save('orphanReport.docx')
            # Adds headers in the 1st row of the table


            e = 0

            child2 = removeAfter(child) #removes everything after the parent tag if there is anything to remove
            # while loop until all the  parentTags has been added to the report


            parents2 = copy.deepcopy(parentTags) # copy of parent tags list
            parents2Copy.extend(parents2)
            childCopy = copy.deepcopy(child2)
            noParent = []
            noParent2 = []
            global orphanChild
            orphanChild = []
            orphanChildParent = []
            parents9000 = []

            parents2 = [s.replace(" ", "") for s in parents2] # gets rid of space
            while parentTags:
                #row = table.add_row().cells # Adding a row and then adding data in it.
                #row[0].text = parentTags[0] # Adds the parentTag to the table
                #report3.add_paragraph(parentTags[0])

                
                
                noParent.append(parentTags[0])


                #for child100 in child2:
                 #   report3.add_paragraph(child100)
                for ch in child2:
                    if "[" not in ch:
                        child2.remove(ch)
                    
                           

                if e < len(fullText2):  #as long as variable e is not higher than the lines in fullText2
                    if fullText2[e] in filtered_LCopy: #filtered_L contains the child tags without a parent tag
                        #report3.add_paragraph(parentTags[0] + " has no child tag")
                        orphanChild.append(parentTags[0])
                        orphanChildren2.append(parentTags[0])
                        #orphanReport.add_paragraph(parentTags[0] + " has no child tag")
                        parentTags.remove(parentTags[0]) # Removes that tag after use
                        noParent2.append(" ")
                        parents9000.append(" ")
                        orphanChildParent.append(" ")
                        #row[1].text = " " # No parent tag, so adds empty string to that cell
                    
                        if child2:
                            if "[" not in child2[0]: # if it is not a parent tag
                                child2.remove(child2[0])  # Removed that tag from the list

                        
                        #if child2:
                         #   child2.remove(child2[0])  # Removed that tag from the list

                        e += 1

                    elif fullText2[e] not in filtered_LCopy:
                        parentTags.remove(parentTags[0]) # Removes that tag after use
                        if child2:
                            #row[1].text = child2[0] #Adds childTag to table
                            
                            if "[" not in child2[0]:
                                child2.remove(child2[0])  # Removed that tag from the list
                                #report3.add_paragraph(child2[0])
                            #report3.add_paragraph(child2[0])
                                        
                            parents9000.append(child2[0])
                            noParent.append(child2[0])
                            child2.remove(child2[0])  # Removed that tag from the list
                            e += 1

            parents9.extend(parents9000)
            orphanChildren2Copy.extend(orphanChildren2)

            # Make sure everything is cleared before the program gets the next document
            child2.clear()
            parentTags.clear()
            child.clear()
            #report3.save('report.docx') #Saves in document "report3"


            global dicts11
            dicts11 = dict(zip(parents2, childCopy)) #creates a dictionary if there is a child tag and parent tag
            dicts.update(dicts)

            noParent = [s.replace(" ", "") for s in noParent]
            #dicts3 = dict(zip(noParent, noParent2)) # dictionary for parent tags without child tags
            orphanChild = [s.replace(" ", "") for s in orphanChild]

            
            #orphanChildren2Copy = copy.deepcopy(orphanChildren2) # copy of orphanChildren2 list

            dicts9000 = dict(zip(orphanChild, orphanChildParent)) # orphan dictionary
            orphanDicts.update(dicts9000)
            OrphanChild2.extend(orphanChild)

            text2 = removeParent(everything) # child tag and text
            # print(text2)
            text3 = removechild(text2)  # only text list
            # print(text3)
            text4 = removeText(text2) # child tags
            # print(text4) #only parent tag list
            #text8 = [s.replace(" ", "") for s in text4]

            parents9000 = [x.strip(' ') for x in parents9000]
            #dicts3 = dict(zip(parents2, childCopy))
            dicts3 = dict(zip(parents2, parents9000))
            dicts10.update(dicts3)
            dicts2 = dict(zip(parents2, text3)) # creates a dictionary with child tags and text
            dicts100 = copy.deepcopy(dicts2)
            
            sorted(dicts2.keys()) # sorts the keys in the dictionary
            dicts2Copy.update(dicts100)




            # for tg in parents2:
            #     if "TBV:" in tg:
            #         TBVReport.add_paragraph(tg)
            # TBVReport.save('TBVReport.docx')
            
        return filepath2, filtered_L, orphanChild
        return parents2, dicts2, dicts10, dicts2Copy, parents2Copy, fullText2, filtered_LCopy, dicts3, orphanDicts, OrphanChild2
        
    except Exception as e:
        # Log an error message
        logging.error('generateReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport(): PASS')


def generateReport2():
    try:

        #print("here is dicts10 before:")
        #print(dicts10)
        global dicts11111 # Will be used for the excel report later for child - parent
        dicts11111 = {}
        dicts11111 = copy.deepcopy(dicts10)

        # counters for Excel report2

        global counter1
        counter1 = 2
        global counter2
        counter2 = 1
        global counter3
        counter3 = 0
        global cell
        cell = 0;
        global cell2
        cell2 = 0;



        #excelReport2.range("A3").value = 'Childless Tags'

        pattern = r'\[([^\]]+)\]'  
        for key in dicts10:
                if type(dicts10[key]) == str:
                    matches2 = re.findall(pattern, (dicts10[key])) 
                if len(matches2) > 1:
                        
                        
                        #print("There is more than one ']' in the input string.", dicts10[key] )
                        dicts10[key] = []
                        parents2 = []
                        for match in matches2:
                            parents2.append(match)
                            
                        for tag in parents2:
                            #parentChild2.setdefault("[PUMP:SRS:1]", ["[PUMP:PRS:0]"]).append("[PUMP:PRS:2]")
                            #parentChild2.setdefault(key, [parentChild2[key]]).append(tag)
                            #parentChild2[key].append(tag)
                            tag = (tag.replace(' ', ''))
                            dicts10[key] += [tag]
                            #parentChild2.setdefault(key, ["[PUMP:PRS:0]"]).append(tag)
                            #parentChild2.setdefault(key, [parentChild2[key]]).append(match)
                            
                else:
                    #print("There is only one ']' in the input string.")
                    #print("")
                    doNothing =+ 1
        
        #print("here is dicts10 after:")
        #print(dicts10)


        #report3.add_paragraph("all parents:") # header for all parents

        parents10 = [] # list of all the parent tag tags
        for value11 in dicts10.values():
            # if the value is a list, extend the parents list with the list
            if isinstance(value11, list):
                parents10.extend(value11)
                
            # if the value is not a list, append the value to the parents list
            else:
                parents10.append(value11)

        # for loop to add the parents to the report
        #for parent in parents10:
         #   report3.add_paragraph(parent)
        

        # create a list of all the keys in the dictionary (all child tags)
        values_list = list(dicts2Copy.keys())

        #  creates a list of all the child tags that are not in the parents list
        global childless
        childless = [] 

                    
        # for loop to check if the child tag is in the parents list
        for element in values_list:
            if "".join(element) not in "".join(parents10):
                childless.append(element)


        # sorts the childless list
        childless.sort()
        print("\n")
        print("childless tag: ")
        # for loop to add the childless tags to the report
        for child0 in childless:
            #childlessReport.add_paragraph(child0) # Adds the childless tags to the report
            print(child0)
        print("\n")

        #childlessReport.save('childless.docx') #Saves in document "childless.docx"

        

        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0
        #print("here is dicts10:", dicts10)

        orphanTagText = removechild(filtered_LCopy)

        #dict(sorted(dicts2Copy.items(), key=lambda item: item[1])) # sorts by value/parentTag, not working at the moment
        #print(parents2Copy)
        #print(dicts10)
        #print(dicts2Copy)
        #print(filtered_LCopy)
        #print(fullText2Copy)
        #print(parents2Copy)
        #print(orphanTagText)
        #print(dicts2Copy)
        #report3.add_paragraph("\n") # Adds a line space from the table
        while m < len(dicts2Copy):
            #print(m)
            #if fullText2Copy[k] not in filtered_LCopy:
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1
                duplicates = []
                for key, value in dicts2Copy.items():
                    
                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        #for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                        #report3.add_paragraph("\n")
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        if str(stringKey2) in dicts10: # if the key is in the dictionary
                            text = dicts10[str(stringKey2)]
                        
                        

                        if isinstance(text, list):
                            #print("it is a list")
                            #report3.add_paragraph("List tags found") # display the parent tag, included brackets
                            
                            for tag in text:
                                #tag = ("[" + tag)
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()
                            
                                #report3.add_paragraph(tag)
                                if (str(tag) in duplicates):
                                    #print("in duplicates")
                                    #print("")
                                    
                                    doNothing += 1
                                    
                                    
                                    

                                else:
                                    parentTag1 = ('['+tag+']')
                                    
                                    cell = str('A'+ str(counter1))
                                    cell2 = str(str(parentTag1))
                                    #excelReport2.range(cell).value = cell2
                                    counter1 += 2 # counter for excel report
                                    counter2 += 1 # counter for excel report

                                    #wb2.save('report2.xlsx') # Saving excel report as 'report2.xlsx'
                                    if "TBV:" in parentTag1:
                                        print("TBV found")
                                       
                                        #TBVReport.add_paragraph(parentTag1)
                                        #TBVReport.save('TBVReport.docx')
                                   

                                    #report3.add_paragraph(parentTag1)
                                    print("parent: ", parentTag1)
                                    tag.strip()
                                    duplicates.append(str(tag))
                                    
                                    

                                    for x in PTags:
                                        #report3.add_paragraph(x)

                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()
                                                                
                                        

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                if "TBV:" in parentTag1:
                                                    print("TBV found")
                                          
                                                print(dicts2Copy[str(keyCheck4)])
                                                #report3.add_paragraph(dicts2Copy[str(keyCheck4)])

                                                
                                            
                                            #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                        else:
                                            print("Requirement text not found")
                                            orphanChildren2Copy.append(str(keyCheck4))
                                            #report3.add_paragraph("Requirement text not found")
                                            #orphanReport.add_paragraph("Requirement text not found")
                                        #print(dicts10[str(key)])
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        #for dicts10[str(stringKey)] in dicts10:
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        for b in PTags:
                                            b = (b.replace(']', ''))
                                            #report3.add_paragraph("I'm b")
                                            #report3.add_paragraph(b)
                                            #report3.add_paragraph("I'm tag")
                                            #report3.add_paragraph(tag)


                                            if b == tag:
                                                i += 1
                                                hx = tag
                                                #report3.add_paragraph("I'm here")
    
                                                keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                

                                                
                                                #for h, v in dicts10.items():
                                                 #   print("hx", hx)
                                                  #  print("v", v)
                                                #report3.add_paragraph("I'm keys")
                                                #report3.add_paragraph(keys)
                                                k += 1
                                                for item in keys: #keys are child tags of hx/the parent tag
                                                    

                                                    if item != "" and item!= " ":
                                                        print("child: ", item)
                                                        print(dicts2Copy[str(item)])
                                                        #report3.add_paragraph(item, style='List Bullet')
                                                        #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                        #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text

                                                        if "TBV:" in parentTag1:
                                                            print("TBV found")
                                                            #TBVReport.add_paragraph(item, style='List Bullet')
                                                            #para2 = TBDReport.add_paragraph(dicts2Copy[str(item)])
                                                            #para2.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                            #TBVReport.save('TBVReport.docx')
                                                        
                                                        
                                                        
                                                        counter2 = counter1 - 1
                                                        cell = str('B'+ str(counter2))
                                                        cell2 = str(item)
                                                        #excelReport2.range(cell).value = cell2
                                                        counter2 += 1
                                                        counter1 += 1


                                                        

                                                #report3.add_paragraph("\n")
                                                print("\n")
                                                counter2 += 1
                                                counter1 += 1
                                                #excelReport2.autofit()



                        else:
                            #print("not a list")    
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()
                            hx10 = text
                            hx10 = hx10.replace('[', '')
                            hx10 = hx10.replace(']', '')
                            #tag.strip()
                            if (str(hx10) in duplicates):
                                #print("in duplicates")
                                #print("")
                                doNothing += 1
                                    

                            else:
                                #parentTag1 = ('['+tag+']')
                                #report3.add_paragraph(parentTag1)
                                #tag.strip()
                                

                                for x in PTags:
                                    
                                    #report3.add_paragraph(str(text))
                                    keyCheck = (x.replace('[', ''))
                                    keyCheck2 = (keyCheck.replace(']', ''))
                                    keyCheck3 = (keyCheck2.replace(']', ''))
                                    keyCheck4 = (keyCheck3.replace(' ', ''))
                                    #report3.add_paragraph(x) # display the parent tag, included brackets
                                
                                    print("parent tag: ", x)
                                    if "TBV:" in x:
                                                            print("TBV found")
                                                            
                                                            #TBVReport.add_paragraph(x)
                                                            #TBVReport.save('TBVReport.docx')

                                    cell = str('A'+ str(counter1))
                                    cell2 = str(str(x))
                                    #excelReport2.range(cell).value = cell2
                                    counter1 += 2 # counter for excel report
                                    counter2 += 1 # counter for excel report

                                    #wb2.save('report2.xlsx') # Saving excel report as 'report2.xlsx'

                                    if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                        if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                            print(dicts2Copy[str(keyCheck4)])
                                            #report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                            if "TBD:" in keyCheck4:
                                                            print("TBD found")
                                                            #TBDReport.add_paragraph(dicts2Copy[str(keyCheck4)])
                                                            #TBDReport.save('TBDReport.docx')

                                        #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                    else:
                                        print("Requirement text not found")
                                        orphanChildren2Copy.append(str(keyCheck4))
                                        #report3.add_paragraph("Requirement text not found")
                                        #orphanReport.add_paragraph("Requirement text not found")
                                    #print(dicts10[str(key)])
                                    #report3.add_paragraph(dicts10[str(stringKey)])
                                    #for dicts10[str(stringKey)] in dicts10:
                                    #report3.add_paragraph(dicts10[str(stringKey)])
                                    for b in PTags:


                                        if b == dicts10[str(stringKey2)]:
                                            i += 1
                                            text.strip()
                                            
                                            hx = str(text)
                                            hx = hx.replace('[', '')
                                            hx = hx.replace(']', '')
                                            
                                            duplicates.append(str(hx))
                                            #report3.add_paragraph(str(hx))
                                            
                                            keys = [h for h, v in dicts10.items() if check_string(hx, v)]
                                            
                                            #for h, v in dicts10.items():
                                             #       print("hx", hx)
                                              #      print("v", v)
                                            # finds all the child tags
                                            #print(keys)
                                            k += 1

                                            #if keys:
                                            #   print("list is not empty")
                                            #else:
                                            #    report3.add_paragraph("It is an orphan tag")

                                            for item in keys: #keys are child tags of hx/the parent tag

                                                if item != "" and item!= " ":
                                                    print("child: ", item)
                                                    print(dicts2Copy[str(item)])
                                                    #report3.add_paragraph(item, style='List Bullet')
                                                    #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                    #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                    #if "TBD:" in item:
                                                     #       TBDReport.add_paragraph(item, style='List Bullet')
                                                      #      para = TBDReport.add_paragraph(dicts2Copy[str(item)])
                                                       #     para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                        #    TBDReport.save('TBDReport.docx')
                                                    #if "TBV:" in item:
                                                    #        TBVReport.add_paragraph(item, style='List Bullet')
                                                     #       para = TBVReport.add_paragraph(dicts2Copy[str(item)])
                                                      #      para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text
                                                       #     TBVReport.save('TBVReport.docx')
                                                    
                                                    
                                                    counter2 = counter1 - 1
                                                    cell = str('B'+ str(counter2))
                                                    cell2 = str(item)
                                                    #excelReport2.range(cell).value = cell2
                                                    counter2 += 1
                                                    counter1 += 1
                                            print("\n")
                                            #report3.add_paragraph("\n")
                                            counter2 += 1
                                            counter1 += 1
                                            #excelReport2.autofit()

                            #report3.add_paragraph("\n") # Adds a line space
                            #print(k)
                            #print(m)
                            #report3.add_paragraph(key, style='List Bullet')
                            #para = report3.add_paragraph(value)
                            #para.paragraph_format.left_indent = Inches(0.25) # adds indentation ot text
                            #stringKey = dicts2Copy[str(key)]
                            #stringKey2 = (stringKey.replace(' ', ''))
                        #if k < len(fullText2Copy):
                    #elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1
                        
                        #report3.add_paragraph("\n")
                        if 1 ==1:
                            for orphantag in orphanChildren2Copy:
                                
                                if orphantag in dicts10:
                                    print("orphantag: ", orphantag)
                                    if orphantag not in duplicates:
                                        kek = 0
                                            #print("orphan parent tag: ", orphantag)
                                            #hx = orphantag
                                            #hx = hx.replace('[', '')
                                            #hx = hx.replace(']', '')
                        
                                            #duplicates.append(str(hx))
                                            #report3.add_paragraph(str(hx))
                                            
                                            #keys = [h for h, v in dicts10.items() if check_string(hx, v)]
                                            #for item in keys: #keys are child tags of hx/the parent tag

                                             #   if item != "" and item!= " ":
                                              #      print("child: ", item)
                                               #     print(dicts2Copy[str(item)])


                            #report3.add_paragraph(parents2Copy[i])
                            #orphanReport.add_paragraph(parents2Copy[i])
                            #paragraph2.add("\n" +parents2Copy[i] + " Is an orphanTag" + "\n")
                            
                            #if "TBV:" in parentTag1:
                            #    TBVReport.add_paragraph(item, style='List Bullet')
                            #    TBVReport.save('TBVReport.docx')
                           # if "TBV:" in parentTag1:
                            #    TBVReport.add_paragraph(item, style='List Bullet')
                            
                            #TBVReport.save('TBVReport.docx')
                            #print(parents2Copy[i])
                            #print(orphanTagText[o])
                            #print("")
                            doNothing += 1
                            #orphanss.append(parents2Copy[i])
                        if o < len(orphanTagText):
                            #report3.add_paragraph(orphanTagText[o])
                            #orphanReport.add_paragraph(orphanTagText[o])
                            #print("")
                            doNothing += 1
                        o += 1
                        if i < len(parents2Copy):
                            #print("")
                            doNothing += 1
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                        #m += 1

                        i += 1
                        #del dicts2Copy[list(dicts2Copy.keys())[0]] # deletes the first item in dicts2Copy

        
        # Description, this function is removing some linees, but not all, maybe need more work
        # We check if the length of the paragraph text is less than 4 using the len() function, 
        # and if so, we remove the paragraph using the _element.clear() method. Finally
        # iterate over all the paragraphs in the document
        #for i in range(len(report3.paragraphs)):
            # check if the paragraph contains less than 4 characters or a string of length less than 4
         #   if len(report3.paragraphs[i].text) < 4:
                # remove the paragraph
          #      report3.paragraphs[i]._element.clear()
                # save the modified document
           #     report3.save('report3.docx')

        

        orphanGenReport()
        #print("dicts10", dicts10)


        
    except Exception as e:
        # Log an error message
        logging.error('generateReport2(): ERROR', e)
    else:
        # Log a success message
        logging.info('generateReport2(): PASS')

    """
        elif not dicts2Copy: # this is for orphan tags
            dict3 = dict(dicts2.items() - dicts3.items())
            for key, value in dicts3.items():
                report3.add_paragraph("\n")
                report3.add_paragraph(key)
                report3.add_paragraph(value)
                report3.add_paragraph(key + " is an orphan tags")
                m += 1
    """

def orphanGenReport():
    duplicates = []
    try:
        # declaring counters
        m = 0
        k = 0
        i = 0
        o = 0
        z = 0

        orphanTagText = removechild(filtered_LCopy)
        while m < len(dicts2Copy):
            #print(m)
            #if fullText2Copy[k] not in filtered_LCopy:
            if z < len(dicts2Copy) and dicts2Copy:
                z += 1

                for key, value in dicts2Copy.items():
                    #orphanReport.add_paragraph("\n")
                    m += 1
                    if k < len(fullText2Copy) and fullText2Copy[k] not in filtered_LCopy:
                        #for key, value in dicts2Copy.items() and key, value in dicts3.items(): #work on this here and try
                        #report3.add_paragraph("\n")
                        stringKey = str(key)
                        stringKey2 = (stringKey.replace(' ', ''))
                        text = dicts10[str(stringKey2)]
                        

                        if isinstance(text, list):
                            #print("it is a list")
                            #print("")
                            doNothing += 1
                            #report3.add_paragraph("List tags found") # display the parent tag, included brackets
                            
                            for tag in text:
                                #tag = ("[" + tag)
                                PTags = tag.split(']')
                                PTags = [s.strip() + ']' for s in PTags]
                                tag.strip()
                                if (str(tag) in duplicates):
                                    #print("in duplicates")
                                    #print("")
                                    doNothing += 1
                                    

                                else:
                                    
                                    #report3.add_paragraph(tag)
                                    #tag.strip()
                                    duplicates.append(str(tag))
                                    
                  
                                    for x in PTags:
                                        #report3.add_paragraph(x)

                                        keyCheck = (x.replace('[', ''))
                                        keyCheck2 = (keyCheck.replace(']', ''))
                                        keyCheck3 = (keyCheck2.replace(']', ''))
                                        keyCheck4 = (keyCheck3.replace(' ', ''))
                                        keyCheck4.split()
                                                                
                                        

                                        if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                            if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                                #report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                    
                                                #print("")
                                                doNothing += 1
                                            #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                        else:
                                            #print("")
                                            doNothing += 1
                                            #report3.add_paragraph("Requirement text not found")
                                            #orphanReport.add_paragraph("Requirement text not found")
                                        #print(dicts10[str(key)])
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        #for dicts10[str(stringKey)] in dicts10:
                                        #report3.add_paragraph(dicts10[str(stringKey)])
                                        for b in PTags:
                                            b = (b.replace(']', ''))
                                            #report3.add_paragraph("I'm b")
                                            #report3.add_paragraph(b)
                                            #report3.add_paragraph("I'm tag")
                                            #report3.add_paragraph(tag)


                                            if b == tag:
                                                i += 1
                                                hx = str(tag)
                                                #report3.add_paragraph("I'm here")
                                                keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                                #for h, v in dicts10.items():
                                                 #   print("hx", hx)
                                                  #  print("v", v)
                                                #report3.add_paragraph("I'm keys")
                                                #report3.add_paragraph(keys)
                                                k += 1
                                                for item in keys: #keys are child tags of hx/the parent tag

                                                    if item != "" and item!= " ":
                                                        #print("")
                                                        doNothing += 1
                                                        #report3.add_paragraph(item, style='List Bullet')
                                                        #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                        #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text





                        else:
                            #print("")
                            doNothing =+ 1
                            PTags = text.split(']')
                            PTags = [s.strip() + ']' for s in PTags]
                            PTags.pop()

                            for x in PTags:
                                keyCheck = (x.replace('[', ''))
                                keyCheck2 = (keyCheck.replace(']', ''))
                                keyCheck3 = (keyCheck2.replace(']', ''))
                                keyCheck4 = (keyCheck3.replace(' ', ''))
                                #report3.add_paragraph(x) # display the parent tag, included brackets

                                if keyCheck4 in dicts2Copy:  # Checks if text of parent tag is found
                                    if dicts2Copy[str(keyCheck4)] != "" and dicts2Copy[str(keyCheck4)] != " ":
                                        #print("")
                                        doNothing += 1
                                        #report3.add_paragraph(dicts2Copy[str(keyCheck4)])
                                    #orphanReport.add_paragraph(dicts2Copy[str(keyCheck4)])

                                else:
                                    #report3.add_paragraph("Requirement text not found")
                                    #print("")
                                    doNothing += 1
                                    #orphanReport.add_paragraph("Requirement text not found")
                                #print(dicts10[str(key)])
                                #report3.add_paragraph(dicts10[str(stringKey)])
                                #for dicts10[str(stringKey)] in dicts10:
                                #report3.add_paragraph(dicts10[str(stringKey)])
                                for b in PTags:

                                    

                                    if b == dicts10[str(stringKey2)]:
                                        i += 1
                                        hx = str(dicts10[str(stringKey2)])
                                        keys = [h for h, v in dicts10.items() if check_string(hx, v)] # finds all the child tags
                                        
                                        #for h, v in dicts10.items():
                                         #           print("hx", hx)
                                          #          print("v", v)
                                        #print(keys)
                                        k += 1
                                        for item in keys: #keys are child tags of hx/the parent tag

                                            if item != "" and item!= " ":
                                                #print("")
                                                doNothing += 1
                                                #report3.add_paragraph(item, style='List Bullet')
                                                #para = report3.add_paragraph(dicts2Copy[str(item)])
                                                #para.paragraph_format.left_indent = Inches(0.25) # adds indentation of text

                            #report3.add_paragraph("\n") # Adds a line space
                            #print(k)
                            #print(m)
                            #report3.add_paragraph(key, style='List Bullet')
                            #para = report3.add_paragraph(value)
                            #para.paragraph_format.left_indent = Inches(0.25) # adds indentation ot text
                            #stringKey = dicts2Copy[str(key)]
                            #stringKey2 = (stringKey.replace(' ', ''))
                        #if k < len(fullText2Copy):
                    #elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                
                        
                    elif k < len(fullText2Copy) and fullText2Copy[k] in filtered_LCopy:
                        k += 1
                        #orphanReport.add_paragraph("\n")
                        if i < len(parents2Copy):
                            orphanss.append(parents2Copy[i]) # adds orphan tags to a list
                            #orphanReport.add_paragraph(parents2Copy[i])
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                            #print(parents2Copy[i])
                            #print(orphanTagText[o])
                            #print("")
                            doNothing += 1
                        if o < len(orphanTagText):
                            #orphanReport.add_paragraph(orphanTagText[o])
                            #print("")
                            doNothing += 1
                            
                        o += 1
                        if i < len(parents2Copy):
                            #print("")
                            doNothing += 1
                            #orphanReport.add_paragraph(parents2Copy[i] + " is an orphan tag")
                        i += 1

        #orphanss.sort() # sorts the list of orphan tags
        #orphanChildren2Copy.sort() # sorts the list of orphan child tags
        
        
        #for orph in orphanss:
        #    orphanReport.add_paragraph(orph)
        #orphanReport.add_paragraph("Orphan Tags2: ")
        print("Orphan Tags: ")
        for orph5 in orphanChildren2Copy:
            #orphanReport.add_paragraph(orph5)
            print(orph5)

        
        return dicts2Copy

    except Exception as e:
        # Log an error message
        logging.error('orphanReport(): ERROR', e)
    else:
        # Log a success message
        logging.info('orphanReport(): PASS')



def removeParent(text): #removes parent tags or child tags
    try:
        childAfter = []
        for line in text:
            childAfter = [i.rsplit('[', 1)[0] for i in text] # removes parent tags
            childAfter = [re.sub("[\(\[].*?[\)\]]", "", e) for e in childAfter]  # removes parent tags that are left
            childAfter = [re.sub("[\{\[].*?[\)\}]", "", e) for e in childAfter]  # removes "pass", "fail", etc.
        return childAfter

    except Exception as e:
        # Log an error message
        logging.error('removeParent(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeParent(): PASS')


def removeText(text6): #this should remove everything before the parent tag
    try:
        childAfter = [s.split(None, 1)[0] if len(s.split(None, 1)) >= 2 else '' for s in text6]
        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeText(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeText(): PASS')


def removeAfter(childtags): #removes everything after the  tag, example "pass"
    try:
        seperator = ']'
        childAfter = [i.rsplit(']', 1)[0] + seperator for i in childtags]

        return childAfter
    except Exception as e:
        # Log an error message
        logging.error('removeAfter(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removeAfter(): PASS')


def removechild(text): #removes child, this one needs fixing
    try:
        mylst = []
        mylst = [s.split(None, 1)[1] if len(s.split(None, 1)) >= 2 else '' for s in text]
        return mylst
    except Exception as e:
        # Log an error message
        logging.error('removechild(): ERROR', exc_info=True)
    else:
        # Log a success message
        logging.info('removechild(): PASS')



def check_string(string1, string2): # checks if a string1 is in string2
    if isinstance(string2, str): 
        string2 = [string2]
    pattern = r'{}(?!\d)'.format(re.escape(string1))
    for s in string2:
        match = re.search(pattern, s)
        if match is not None:
            return True
    return False






